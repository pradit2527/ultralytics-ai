"""
AIDC Tech Video Processor — เว็บ UI ประมวลผลไฟล์วิดีโอด้วย Ultralytics YOLO
รัน:  D:\\yoloe\\venv\\Scripts\\python.exe D:\\yoloe\\app.py
แล้วเปิดเบราว์เซอร์ตามลิงก์ที่ขึ้น (ปกติ http://127.0.0.1:7860)
"""

import os
import shutil
import subprocess
import tempfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

import cv2
import gradio as gr
import pandas as pd
import torch
from ultralytics import YOLO, YOLOWorld

# โมเดล open-vocabulary สำหรับตรวจจับวัตถุ "ตามคำที่พิมพ์" (เช่น tree)
WORLD_MODEL = "yolov8s-world.pt"
# โมเดล segmentation (วาด mask ขอบเขตวัตถุ) — รุ่นเดียวกับ yolo26n ดาวน์โหลดอัตโนมัติ
SEG_MODEL = "yolo26n-seg.pt"

# รายการวัตถุสำเร็จรูป แบ่งเป็นหมวดหมู่: หมวด → { ป้ายไทย → คำสั่งอังกฤษ }
CATEGORIES = {
    "คนและสัตว์": {
        "👤 คน": "person",
        "🐾 สัตว์": "animal",
        "🐶 สุนัข": "dog",
        "🐦 นก": "bird",
    },
    "ธรรมชาติและภูมิประเทศ": {
        "🌳 ต้นไม้": "tree",
        "🪵 ท่อนไม้ใหญ่": "wood log",
        "🌿 พุ่มไม้": "bush",
        "🌊 ลำธาร/แม่น้ำ": "river",
        "⛰️ ภูเขา": "mountain",
        "🪨 ก้อนหิน": "rock",
    },
    "ไฟและควัน": {
        "🔥 กองไฟ/ไฟ": "fire",
        "💨 ควันไฟ": "smoke",
    },
    "ยานพาหนะและสิ่งก่อสร้าง": {
        "🚗 รถยนต์": "car",
        "🏠 อาคาร/สิ่งปลูกสร้าง": "building",
        "🛣️ ถนน/ทางเดิน": "road",
        "🚧 ป้าย/เครื่องหมาย": "sign",
    },
}
# รวมทุกหมวดเป็น dict เดียว + แผนที่ย้อนกลับ (อังกฤษ → ไทย) ไว้แปลชื่อในตาราง
PRESETS = {th: en for items in CATEGORIES.values() for th, en in items.items()}
PROMPT_TO_TH = {en: th for th, en in PRESETS.items()}

# ชุดสำเร็จรูป (ปุ่มกดเลือกทีเดียวหลายอย่าง)
QUICK_SETS = {
    "🌲 ป่าไม้/ธรรมชาติ": ["🌳 ต้นไม้", "🪵 ท่อนไม้ใหญ่", "🌿 พุ่มไม้",
                          "🌊 ลำธาร/แม่น้ำ", "⛰️ ภูเขา", "🪨 ก้อนหิน"],
    "🔥 ไฟป่า/ควัน": ["🔥 กองไฟ/ไฟ", "💨 ควันไฟ", "🌳 ต้นไม้", "👤 คน"],
    "👥 คน/สัตว์": ["👤 คน", "🐾 สัตว์", "🐶 สุนัข", "🐦 นก"],
}

ROOT = Path(r"D:\yoloe")
# ทำงานในโฟลเดอร์ D:\yoloe เสมอ เพื่อให้โมเดลที่ดาวน์โหลดอัตโนมัติ (เช่น
# yolov8s-world.pt, CLIP) ลงที่นี่ ไม่ไปเลอะโฟลเดอร์อื่นที่สั่งรันแอป
os.chdir(ROOT)
HAS_CUDA = torch.cuda.is_available()
GPU_NAME = torch.cuda.get_device_name(0) if HAS_CUDA else "CPU only"

# แคชโมเดลที่โหลดแล้ว เพื่อไม่ต้องโหลดซ้ำทุกครั้ง
_MODEL_CACHE: dict[str, YOLO] = {}


def find_models() -> list[str]:
    """หาไฟล์ .pt ทั้งหมดในโฟลเดอร์ D:\\yoloe"""
    models = sorted(str(p) for p in ROOT.glob("*.pt"))
    return models or ["yolo26n.pt"]


def load_animals() -> list[str]:
    """โหลดรายชื่อสัตว์จากไฟล์ animals.txt (ใช้เป็นตัวเลือกตรวจจับแบบละเอียด)"""
    p = ROOT / "animals.txt"
    if p.exists():
        return [ln.strip() for ln in p.read_text(encoding="utf-8").splitlines()
                if ln.strip()]
    return []


ANIMALS = load_animals()

# จัดกลุ่มสัตว์ตามตัวอักษรขึ้นต้น เพื่อทำตัวกรองให้รายการไม่ล้นในดรอปดาวน์เดียว
ANIMAL_ALL = "__ALL__"


def _group_animals_by_letter(names: list[str]) -> dict[str, list[str]]:
    groups: dict[str, list[str]] = {}
    for n in names:
        first = n[0].upper() if n else "#"
        if not ("A" <= first <= "Z"):
            first = "#"
        groups.setdefault(first, []).append(n)
    return groups


ANIMAL_GROUPS = _group_animals_by_letter(ANIMALS)
# ตัวเลือกหมวด (label, value) — โชว์จำนวนในวงเล็บ
ANIMAL_LETTER_CHOICES = [(f"⭐ ทั้งหมด ({len(ANIMALS)})", ANIMAL_ALL)] + [
    (f"{L}  ({len(ANIMAL_GROUPS[L])})", L) for L in sorted(ANIMAL_GROUPS)
]


def filter_animals(group, selected):
    """ย่อรายการสัตว์ตามหมวดตัวอักษร — แต่คงรายการที่เลือกไว้แล้วเสมอ"""
    base = ANIMALS if group in (None, ANIMAL_ALL) else ANIMAL_GROUPS.get(group, [])
    extra = [x for x in (selected or []) if x not in base]
    return gr.update(choices=base + extra)


def get_model(model_path: str) -> YOLO:
    if model_path not in _MODEL_CACHE:
        _MODEL_CACHE[model_path] = YOLO(model_path)
    return _MODEL_CACHE[model_path]


_WORLD_CLASSES: list[str] = []  # คลาสล่าสุดที่ตั้งไว้ (กันการตั้งซ้ำโดยไม่จำเป็น)


def get_world_model(classes: list[str]) -> YOLOWorld:
    """โมเดล open-vocabulary — ตั้งคลาสตามคำที่ผู้ใช้พิมพ์ (เช่น tree, dog)"""
    if WORLD_MODEL not in _MODEL_CACHE:
        _MODEL_CACHE[WORLD_MODEL] = YOLOWorld(WORLD_MODEL)
    model = _MODEL_CACHE[WORLD_MODEL]
    if list(classes) != _WORLD_CLASSES:
        # set_classes() ต้องเข้ารหัสข้อความด้วย CLIP ก่อน — ต้องทำบน CPU เสมอ
        # ไม่งั้นจะเจอบั๊ก "tensors on different devices": CLIP ถูกย้ายไป GPU
        # จากการ predict รอบก่อน แต่ token ข้อความถูกสร้างบน CPU แล้วชนกัน
        # วิธีแก้: ย้ายโมเดลกลับ CPU + ล้าง CLIP cache ให้สร้างใหม่บน CPU
        # (ตอน predict ระบบจะย้าย text features ไป GPU ให้เองอยู่แล้ว)
        model.to("cpu")
        try:
            model.model.clip_model = None
        except Exception:
            pass
        model.set_classes(classes)
        _WORLD_CLASSES[:] = list(classes)
    return model


def parse_classes(text: str) -> list[str]:
    """แยกคำที่คั่นด้วย , หรือขึ้นบรรทัดใหม่ → ลิสต์คลาส"""
    return [c.strip() for c in (text or "").replace("\n", ",").split(",") if c.strip()]


def _ffmpeg_exe():
    """หา ffmpeg: ใช้ของระบบก่อน ไม่มีก็ใช้ตัวที่มากับ imageio-ffmpeg"""
    exe = shutil.which("ffmpeg")
    if exe:
        return exe
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except Exception:
        return None


def _reencode_h264(src: str) -> str:
    """แปลงเป็น H.264 ให้เบราว์เซอร์เล่นในหน้าเว็บได้ (เล่นปุ่ม play ได้ทันที)"""
    exe = _ffmpeg_exe()
    if not exe:
        return src
    dst = src.replace(".mp4", "_h264.mp4")
    try:
        subprocess.run(
            [exe, "-y", "-i", src, "-c:v", "libx264", "-pix_fmt", "yuv420p",
             "-movflags", "+faststart", "-an", dst],
            check=True, capture_output=True,
        )
        return dst
    except Exception:
        return src


def counts_to_df(counts, tracking):
    """แปลงลิสต์ counts → DataFrame (เพิ่มคอลัมน์ ‘นับตัวจริง’ เมื่อเปิด tracking)"""
    base_cols = ["วัตถุ (class)", "ตรวจจับรวม (ทุกเฟรม)", "สูงสุดต่อเฟรม"]
    cols = base_cols + (["นับตัวจริง (ไม่ซ้ำ)"] if tracking else [])
    rows = []
    for c in (counts or []):
        row = {base_cols[0]: c["name"], base_cols[1]: c["total"], base_cols[2]: c["peak"]}
        if tracking:
            row["นับตัวจริง (ไม่ซ้ำ)"] = c.get("unique") if c.get("unique") is not None else "-"
        rows.append(row)
    return pd.DataFrame(rows) if rows else pd.DataFrame(columns=cols)


def process_video(video_path, model_path, preset_labels, animal_labels, custom_classes,
                  segment, track, conf, use_gpu, frame_stride, progress=gr.Progress()):
    if not video_path:
        raise gr.Error("กรุณาอัปโหลดไฟล์วิดีโอก่อน")

    device = 0 if (use_gpu and HAS_CUDA) else "cpu"
    seg_note = ""
    # รวม: หมวดที่ติ๊ก + สัตว์ที่เลือกจากรายการละเอียด + ที่พิมพ์เอง (ตัดซ้ำ คงลำดับ)
    chosen = [PRESETS[l] for l in (preset_labels or []) if l in PRESETS]
    chosen += list(animal_labels or [])          # ชื่อสัตว์เป็นภาษาอังกฤษอยู่แล้ว
    chosen = list(dict.fromkeys(chosen + parse_classes(custom_classes)))
    if chosen:
        progress(0, desc=f"กำลังโหลดโมเดล AI สำหรับ: {', '.join(chosen[:8])}"
                         + (" ..." if len(chosen) > 8 else ""))
        model = get_world_model(chosen)   # ตรวจจับตามรายการที่เลือก/พิมพ์
        mode_label = f"AI ตรวจจับเอง ({len(chosen)} ชนิด)"
        if segment:
            # YOLO-World ให้ได้แค่กรอบ ไม่มี mask — แจ้งเตือนแล้วใช้กรอบตามปกติ
            seg_note = " | ⚠️ โหมด mask ใช้ได้เฉพาะตรวจจับ 80 ชนิดมาตรฐาน (ไม่รวมที่เลือกเอง)"
    elif segment:
        progress(0, desc="กำลังโหลดโมเดล segmentation ...")
        model = get_model(SEG_MODEL)      # โมเดล seg วาด mask ขอบเขตวัตถุ
        mode_label = f"{SEG_MODEL} (segmentation)"
    else:
        model = get_model(model_path)     # โมเดลปกติ (80 คลาส COCO)
        mode_label = Path(model_path).name
    stride = max(1, int(frame_stride))

    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise gr.Error("เปิดไฟล์วิดีโอไม่ได้ — ไฟล์อาจเสียหรือไม่รองรับ")

    fps = cap.get(cv2.CAP_PROP_FPS) or 25
    width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) or 0
    out_fps = max(1.0, fps / stride)

    out_path = os.path.join(tempfile.mkdtemp(), "result.mp4")
    writer = cv2.VideoWriter(out_path, cv2.VideoWriter_fourcc(*"mp4v"),
                             out_fps, (width, height))

    total_counts: Counter = Counter()       # รวมจำนวนตรวจจับสะสมต่อคลาส
    peak_counts: dict[str, int] = defaultdict(int)  # จำนวนสูงสุดที่เจอพร้อมกันในเฟรมเดียว
    unique_ids: dict[str, set] = defaultdict(set)   # ID ที่ไม่ซ้ำต่อคลาส (โหมด tracking)
    processed = 0
    idx = 0
    best_n = -1                              # เฟรมที่เจอวัตถุมากสุด (ไว้ส่งให้ Claude)
    best_frame = None
    track_on = bool(track)
    track_active = False                     # เปิดใช้ tracking ได้จริงหรือไม่

    try:
        while True:
            ok, frame = cap.read()
            if not ok:
                break
            if idx % stride == 0:
                if track_on:
                    # นับตัวจริงด้วย tracker (ByteTrack) — persist=False เฉพาะเฟรมแรกเพื่อรีเซ็ต
                    try:
                        res = model.track(frame, conf=float(conf), device=device,
                                          persist=(processed > 0), tracker="bytetrack.yaml",
                                          verbose=False)[0]
                        track_active = True
                    except Exception:
                        track_on = False      # โมเดลนี้ track ไม่ได้ → ถอยไปใช้ detect
                        res = model.predict(frame, conf=float(conf), device=device,
                                            verbose=False)[0]
                else:
                    res = model.predict(frame, conf=float(conf), device=device, verbose=False)[0]

                plotted = res.plot()          # เฟรมที่วาดกรอบ/mask (+ID) แล้ว (BGR)
                frame_classes = Counter()
                for b in res.boxes:
                    name = model.names[int(b.cls)]
                    name = PROMPT_TO_TH.get(name, name)  # แปลเป็นป้ายไทยถ้ามี
                    frame_classes[name] += 1
                    if track_active and b.id is not None:
                        unique_ids[name].add(int(b.id))
                total_counts.update(frame_classes)
                for name, c in frame_classes.items():
                    peak_counts[name] = max(peak_counts[name], c)

                if len(res.boxes) > best_n:    # จำเฟรมที่เด่นสุดไว้
                    best_n = len(res.boxes)
                    best_frame = plotted.copy()

                writer.write(plotted)
                processed += 1

                if total:
                    progress(min(idx / total, 0.99),
                             desc=f"กำลังประมวลผล เฟรม {idx}/{total}")
            idx += 1
    finally:
        cap.release()
        writer.release()

    # บันทึกเฟรมเด่นเป็น JPEG ไว้ให้ Claude วิเคราะห์ภายหลัง
    keyframe_path = None
    if best_frame is not None:
        keyframe_path = os.path.join(tempfile.mkdtemp(), "keyframe.jpg")
        cv2.imwrite(keyframe_path, best_frame)

    progress(0.99, desc="กำลังเข้ารหัสวิดีโอผลลัพธ์...")
    final_path = _reencode_h264(out_path)

    counts = [
        {"name": n, "total": total_counts[n], "peak": peak_counts[n],
         "unique": (len(unique_ids[n]) if track_active else None)}
        for n in sorted(total_counts, key=lambda n: -total_counts[n])
    ]
    df = counts_to_df(counts, track_active)

    track_note = ""
    if track_active:
        track_note = f" | นับตัวจริง {sum(len(s) for s in unique_ids.values())} ตัว"
    summary = (
        f"✅ เสร็จสิ้น | โมเดล: `{mode_label}` | "
        f"ประมวลผล {processed} เฟรม (จาก {total or '?'}) | "
        f"อุปกรณ์: {'GPU — ' + GPU_NAME if device == 0 else 'CPU'} | "
        f"พบวัตถุ {len(total_counts)} ชนิด{track_note}{seg_note}"
    )
    # ข้อมูลสำหรับรายงาน/ส่งออก (ใช้ตอนกดปุ่มรายงาน, Word, Excel)
    report_data = {
        "mode": mode_label,
        "processed": processed,
        "total_frames": total,
        "fps": round(fps, 1),
        "duration_sec": round((total / fps), 1) if total and fps else None,
        "counts": counts,
        "tracking": track_active,
        "keyframe": keyframe_path,
        "video": final_path,
    }
    return final_path, df, summary, report_data


# ── รายงานประเมินความเสี่ยงด้วย Claude AI ──────────────────────────────
RISK_SYSTEM_PROMPT = (
    "คุณเป็นผู้ช่วยวิเคราะห์ภาพจากกล้องสำหรับหน่วยงานราชการ หน้าที่คือประเมิน "
    "ความเสี่ยงและความผิดปกติ จากผลการตรวจจับวัตถุด้วย AI (YOLO) ร่วมกับภาพเฟรม "
    "ตัวอย่างที่แนบมา เขียนรายงานเป็นภาษาไทยที่เป็นทางการและกระชับ ประกอบด้วย: "
    "(1) ระดับความเสี่ยงโดยรวม — ระบุชัดเจนว่า ต่ำ / ปานกลาง / สูง "
    "(2) สิ่งที่ตรวจพบและการตีความ (3) ความผิดปกติหรือจุดที่ควรเฝ้าระวัง "
    "(4) ข้อเสนอแนะ. ใช้เฉพาะข้อมูลที่เห็นจริงในภาพและตัวเลขที่ให้มา "
    "ห้ามกุข้อมูลที่ไม่ปรากฏ หากข้อมูลไม่พอ ให้ระบุข้อจำกัดอย่างตรงไปตรงมา"
)


def _build_report_prompt(report_data) -> str:
    lines = [f"โหมดการตรวจจับ: {report_data['mode']}"]
    if report_data.get("duration_sec"):
        lines.append(f"ความยาววิดีโอโดยประมาณ: {report_data['duration_sec']} วินาที "
                     f"({report_data['fps']} fps)")
    lines.append(f"จำนวนเฟรมที่ประมวลผล: {report_data['processed']}")
    lines.append("\nสรุปวัตถุที่ตรวจพบ (เรียงตามจำนวนรวม):")
    if report_data["counts"]:
        for c in report_data["counts"]:
            lines.append(f"- {c['name']}: ตรวจจับรวมทุกเฟรม {c['total']} ครั้ง, "
                         f"สูงสุดต่อเฟรม {c['peak']}")
    else:
        lines.append("- ไม่พบวัตถุที่ตรวจจับได้")
    lines.append("\nภาพที่แนบคือเฟรมที่ตรวจพบวัตถุมากที่สุด (วาดกรอบ/mask แล้ว) "
                 "โปรดประเมินความเสี่ยงและความผิดปกติพร้อมข้อเสนอแนะ")
    return "\n".join(lines)


def generate_ai_report(report_data, progress=gr.Progress()):
    if not report_data:
        return "⚠️ กรุณากด **เริ่มประมวลผล** วิดีโอก่อน แล้วจึงสร้างรายงาน"
    if not os.environ.get("ANTHROPIC_API_KEY"):
        return ("### ⚠️ ยังไม่ได้ตั้งค่า ANTHROPIC_API_KEY\n\n"
                "1. รับคีย์ที่ https://console.anthropic.com\n"
                "2. เปิด PowerShell พิมพ์ (ใส่คีย์ของคุณ):\n"
                "```\nsetx ANTHROPIC_API_KEY \"sk-ant-...\"\n```\n"
                "3. ปิดแล้วเปิดแอปใหม่ (`run_video_ui.bat`)")
    try:
        import anthropic
        import base64
    except Exception:
        return "⚠️ ไม่พบไลบรารี anthropic — ติดตั้งด้วย `pip install anthropic`"

    progress(0.3, desc="กำลังส่งให้ Claude วิเคราะห์ ...")
    content = []
    kf = report_data.get("keyframe")
    if kf and os.path.exists(kf):
        with open(kf, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode()
        content.append({"type": "image", "source": {
            "type": "base64", "media_type": "image/jpeg", "data": img_b64}})
    content.append({"type": "text", "text": _build_report_prompt(report_data)})

    try:
        client = anthropic.Anthropic()
        msg = client.messages.create(
            model="claude-opus-4-8",
            max_tokens=3000,
            thinking={"type": "adaptive"},
            system=RISK_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": content}],
        )
    except anthropic.AuthenticationError:
        return "⚠️ ANTHROPIC_API_KEY ไม่ถูกต้อง — โปรดตรวจสอบคีย์อีกครั้ง"
    except anthropic.APIError as e:
        return f"⚠️ เรียก Claude ไม่สำเร็จ: {e}"
    except Exception as e:
        return f"⚠️ เกิดข้อผิดพลาด: {e}"

    text = "".join(b.text for b in msg.content if b.type == "text")
    return "## 📝 รายงานประเมินความเสี่ยง (วิเคราะห์โดย Claude AI)\n\n" + text


# ── หน้า "รายงานผล" — รวบรวมผลการวิเคราะห์ล่าสุดเป็นเอกสารทางการ ──────
_TH_MONTHS = ["", "ม.ค.", "ก.พ.", "มี.ค.", "เม.ย.", "พ.ค.", "มิ.ย.",
              "ก.ค.", "ส.ค.", "ก.ย.", "ต.ค.", "พ.ย.", "ธ.ค."]


def _thai_now() -> str:
    n = datetime.now()
    return f"{n.day} {_TH_MONTHS[n.month]} {n.year + 543} เวลา {n:%H:%M} น."


def build_report_page(report_data, ai_text):
    """คืนค่า (HTML หัว+KPI, ภาพเฟรมเด่น, DataFrame ตาราง, Markdown รายงาน AI)"""
    empty_df = counts_to_df([], False)
    if not report_data:
        html = ('<div class="status-box">ยังไม่มีผลการวิเคราะห์ — กรุณาไปที่แท็บ '
                '<b>“วิเคราะห์วิดีโอ”</b> อัปโหลดวิดีโอและกด “เริ่มประมวลผล” ก่อน '
                'แล้วจึงกลับมาที่หน้านี้</div>')
        return html, None, empty_df, ""

    counts = report_data["counts"]
    tracking = report_data.get("tracking")
    total_det = sum(c["total"] for c in counts)
    peak_all = max((c["peak"] for c in counts), default=0)
    unique_all = sum((c.get("unique") or 0) for c in counts) if tracking else None
    dur = report_data.get("duration_sec")

    def kpi(value, label):
        return (f'<div class="kpi"><div class="kpi-v">{value}</div>'
                f'<div class="kpi-l">{label}</div></div>')

    cards = [kpi(len(counts), "ชนิดวัตถุที่พบ"),
             kpi(f"{total_det:,}", "การตรวจจับรวม (ทุกเฟรม)")]
    if tracking:
        cards.append(kpi(unique_all, "นับตัวจริงรวม (ไม่ซ้ำ)"))
    cards.append(kpi(peak_all, "สูงสุดพร้อมกัน/เฟรม"))
    cards.append(kpi(report_data["processed"], "เฟรมที่ประมวลผล"))

    meta = (
        f'<div class="rpt-meta"><span><b>วันที่ออกรายงาน:</b> {_thai_now()}</span>'
        f'<span><b>โหมดตรวจจับ:</b> {report_data["mode"]}</span>'
        + (f'<span><b>ความยาววิดีโอ:</b> {dur} วินาที ({report_data["fps"]} fps)</span>'
           if dur else "")
        + '</div>'
    )
    html = (
        '<div class="rpt-doc">'
        '<div class="rpt-head"><div class="rpt-emblem">🛡️</div>'
        '<div><div class="rpt-org">AIDC TECH</div>'
        '<div class="rpt-title">รายงานผลการวิเคราะห์วิดีโอด้วยปัญญาประดิษฐ์</div></div></div>'
        + meta + f'<div class="kpi-row">{"".join(cards)}</div></div>'
    )

    df = counts_to_df(counts, tracking)
    keyframe = report_data.get("keyframe")
    keyframe = keyframe if (keyframe and os.path.exists(keyframe)) else None
    ai_md = ai_text or ("> ยังไม่ได้สร้างรายงานประเมินความเสี่ยงด้วย AI — "
                        "กดปุ่ม “สร้างรายงานประเมินความเสี่ยง (Claude AI)” "
                        "ในแท็บวิเคราะห์วิดีโอ")
    return html, keyframe, df, ai_md


# ── ส่งออกเอกสาร: Word / Excel / วิดีโอ ────────────────────────────────
def export_excel(report_data):
    if not report_data:
        raise gr.Error("ยังไม่มีผลการวิเคราะห์ — กรุณาประมวลผลก่อน")
    df = counts_to_df(report_data["counts"], report_data.get("tracking"))
    path = os.path.join(tempfile.mkdtemp(), "AIDC_detection_report.xlsx")
    with pd.ExcelWriter(path, engine="openpyxl") as xw:
        df.to_excel(xw, index=False, sheet_name="สรุปการตรวจจับ")
    return path


def export_word(report_data, ai_text):
    if not report_data:
        raise gr.Error("ยังไม่มีผลการวิเคราะห์ — กรุณาประมวลผลก่อน")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Tahoma"          # ฟอนต์รองรับภาษาไทย
    style.font.size = Pt(14)
    style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), "Tahoma")

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = org.add_run("AIDC TECH"); r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0xC9, 0xA1, 0x4A)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("รายงานผลการวิเคราะห์วิดีโอด้วยปัญญาประดิษฐ์")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x0E, 0x2A, 0x5E)

    for line in [f"วันที่ออกรายงาน: {_thai_now()}",
                 f"โหมดตรวจจับ: {report_data['mode']}",
                 (f"ความยาววิดีโอ: {report_data['duration_sec']} วินาที "
                  f"({report_data['fps']} fps)" if report_data.get('duration_sec') else None),
                 f"จำนวนเฟรมที่ประมวลผล: {report_data['processed']}"]:
        if line:
            doc.add_paragraph(line)

    doc.add_paragraph()
    h = doc.add_paragraph(); r = h.add_run("สรุปการตรวจจับ"); r.bold = True; r.font.size = Pt(15)

    counts = report_data["counts"]
    tracking = report_data.get("tracking")
    headers = ["วัตถุ", "ตรวจจับรวม (ทุกเฟรม)", "สูงสุดต่อเฟรม"] + \
              (["นับตัวจริง (ไม่ซ้ำ)"] if tracking else [])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = htext
        for p in cell.paragraphs:
            for rr in p.runs:
                rr.bold = True
    for c in counts:
        cells = table.add_row().cells
        cells[0].text = str(c["name"])
        cells[1].text = str(c["total"])
        cells[2].text = str(c["peak"])
        if tracking:
            cells[3].text = str(c.get("unique") if c.get("unique") is not None else "-")

    kf = report_data.get("keyframe")
    if kf and os.path.exists(kf):
        doc.add_paragraph()
        h = doc.add_paragraph(); r = h.add_run("ภาพเฟรมตัวอย่าง (ตรวจพบวัตถุมากที่สุด)")
        r.bold = True; r.font.size = Pt(15)
        try:
            doc.add_picture(kf, width=Inches(6))
        except Exception:
            pass

    if ai_text:
        doc.add_paragraph()
        h = doc.add_paragraph(); r = h.add_run("รายงานประเมินความเสี่ยง (วิเคราะห์โดย AI)")
        r.bold = True; r.font.size = Pt(15)
        for ln in ai_text.splitlines():
            ln = ln.lstrip("#").replace("**", "").strip()
            if ln:
                doc.add_paragraph(ln)

    doc.add_paragraph()
    foot = doc.add_paragraph("สงวนลิขสิทธิ์ © 2569 AIDC Tech · เพื่อการใช้งานภายในหน่วยงาน")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for rr in foot.runs:
        rr.font.size = Pt(10); rr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    path = os.path.join(tempfile.mkdtemp(), "AIDC_report.docx")
    doc.save(path)
    return path


def get_result_video(report_data):
    v = (report_data or {}).get("video")
    if not v or not os.path.exists(v):
        raise gr.Error("ยังไม่มีวิดีโอผลลัพธ์ — กรุณาประมวลผลก่อน")
    return v


THEME = gr.themes.Soft(
    primary_hue="indigo",
    secondary_hue="violet",
    neutral_hue="slate",
    font=[gr.themes.GoogleFont("Plus Jakarta Sans"),
          gr.themes.GoogleFont("Sarabun"), "system-ui", "sans-serif"],
    radius_size="lg",
)

_dot = "#10b981" if HAS_CUDA else "#94a3b8"
_badge = ("พร้อมใช้งาน · GPU" if HAS_CUDA else "พร้อมใช้งาน · CPU")

CSS = """
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@500;600;700;800&family=Kanit:wght@500;600;700&display=swap');
:root {
  --brand:#6366f1; --brand2:#8b5cf6; --brand3:#a855f7;
  --brand-ink:#4f46e5; --brand-deep:#312e81; --brand-night:#1e1b4b;
  --cyan:#22d3ee; --ok:#10b981; --gold:#f59e0b;
  --ink:#0f172a; --muted:#5b6478; --faint:#94a3b8;
  --line:#e8eaf3; --line2:#eef0f8; --card:#ffffff;
  --tint:#f3f4fd; --tint2:#eef1fe;
  --shadow:30,27,75; --brand-rgb:99,102,241;
}
.gradio-container {max-width:1520px !important; margin:auto !important;
  padding:18px 16px 0 !important; color:var(--ink) !important;
  background:
    radial-gradient(900px 420px at 88% -8%, rgba(139,92,246,.14), transparent 60%),
    radial-gradient(820px 420px at 6% 4%, rgba(99,102,241,.13), transparent 58%),
    linear-gradient(180deg,#eef1fb 0%, #e9ecf8 100%) !important;}
footer {display:none !important;}
.gradio-container * {font-feature-settings:"liga" 1; -webkit-font-smoothing:antialiased;}
h1,h2,h3,.section-title,.pg-title,.rpt-title,.kpi-v,.ft-h,.app-masthead .org,
.app-masthead h1,.mh-ver b,.gov-tabs button[role="tab"]{
  font-family:'Plus Jakarta Sans','Kanit','Sarabun',system-ui,sans-serif !important;}

::selection {background:rgba(139,92,246,.22);}
.gradio-container ::-webkit-scrollbar {width:11px; height:11px;}
.gradio-container ::-webkit-scrollbar-thumb {
  background:linear-gradient(180deg,#c7cbf0,#b9bdec); border-radius:10px;
  border:3px solid transparent; background-clip:padding-box;}
.gradio-container ::-webkit-scrollbar-thumb:hover {background:var(--brand); background-clip:padding-box;}

@keyframes fadeUp {from{opacity:0; transform:translateY(12px);} to{opacity:1; transform:none;}}
@keyframes pulse {0%,100%{box-shadow:0 0 0 0 rgba(16,185,129,.5);} 50%{box-shadow:0 0 0 6px rgba(16,185,129,0);}}
@keyframes drift {0%{transform:translate(0,0) scale(1);} 50%{transform:translate(22px,-16px) scale(1.12);} 100%{transform:translate(0,0) scale(1);}}
.card,.rpt-doc,.app-footer,.app-masthead {animation:fadeUp .45s cubic-bezier(.22,.61,.36,1) both;}

/* ── แถบ utility ด้านบน ───────────────────────────────── */
.app-ubar {display:flex; align-items:center; justify-content:space-between;
  background:linear-gradient(90deg,var(--brand-night),#241f57); color:#c9cdf2;
  font-size:11.5px; letter-spacing:.5px; padding:8px 20px; font-weight:500;
  border-radius:20px 20px 0 0; border-bottom:1px solid rgba(255,255,255,.07);}
.app-ubar .ub-right {display:inline-flex; align-items:center; gap:9px; color:#eef0ff; font-weight:600;}
.app-ubar .ub-dot {width:9px; height:9px; border-radius:50%; background:%DOT%; animation:pulse 2.2s infinite;}

/* ── masthead (เฮดเดอร์หลัก + แสง aurora) ───────────────── */
.app-masthead {position:relative; overflow:hidden; display:flex; align-items:center; gap:22px;
  color:#fff; padding:30px 32px; border-radius:0 0 24px 24px;
  background:linear-gradient(125deg,#0f1437 0%, var(--brand-night) 52%, #3b2e86 100%);
  box-shadow:0 26px 50px -24px rgba(49,46,129,.7);}
.app-masthead .mh-aurora {position:absolute; border-radius:50%; filter:blur(46px); opacity:.62;
  pointer-events:none; animation:drift 13s ease-in-out infinite;}
.app-masthead .a1 {width:280px; height:280px; left:-40px; top:-120px;
  background:radial-gradient(circle,#8b5cf6,transparent 68%);}
.app-masthead .a2 {width:340px; height:340px; right:-60px; bottom:-160px; animation-delay:-4s;
  background:radial-gradient(circle,#22d3ee,transparent 68%);}
.app-masthead .a3 {width:240px; height:240px; left:42%; top:-150px; animation-delay:-8s; opacity:.5;
  background:radial-gradient(circle,#6366f1,transparent 68%);}
.app-masthead::after{content:""; position:absolute; inset:0; pointer-events:none;
  background-image:linear-gradient(rgba(255,255,255,.045) 1px,transparent 1px),
                   linear-gradient(90deg,rgba(255,255,255,.045) 1px,transparent 1px);
  background-size:30px 30px; mask-image:radial-gradient(120% 120% at 50% 0%,#000 35%,transparent 75%);}
.app-masthead > * {position:relative; z-index:1;}
.app-masthead .mh-emblem {width:66px; height:66px; flex:0 0 66px; border-radius:20px;
  background:linear-gradient(160deg,rgba(255,255,255,.22),rgba(255,255,255,.06));
  border:1px solid rgba(196,181,253,.55); display:flex; align-items:center; justify-content:center;
  font-size:33px; backdrop-filter:blur(6px);
  box-shadow:0 10px 26px rgba(0,0,0,.28), inset 0 1px 0 rgba(255,255,255,.3);}
.app-masthead .org {font-size:12px; letter-spacing:5px; font-weight:800;
  background:linear-gradient(90deg,#e9d5ff,#a5b4fc); -webkit-background-clip:text;
  background-clip:text; -webkit-text-fill-color:transparent;}
.app-masthead h1 {margin:6px 0 0; font-size:24px; font-weight:800; color:#fff; line-height:1.32;
  letter-spacing:.2px; text-shadow:0 2px 14px rgba(0,0,0,.3);}
.app-masthead .sub {margin:6px 0 0; font-size:12.5px; color:#c2c9ef; letter-spacing:.2px;}
.app-masthead .mh-ver {margin-left:auto; text-align:right; font-size:12px; color:#c7cdf0;
  border-left:1px solid rgba(255,255,255,.16); padding-left:22px; line-height:1.9;}
.app-masthead .mh-ver b {color:#fff; font-size:14px;}

/* ── แถบชื่อหน้า / breadcrumb ─────────────────────────── */
.page-head {position:relative; overflow:hidden; display:flex; align-items:center;
  justify-content:space-between; margin:18px 0 6px; flex-wrap:wrap; gap:8px;
  background:linear-gradient(120deg,#fff,#f6f7fe); padding:16px 22px;
  border:1px solid var(--line); border-radius:16px;
  box-shadow:0 12px 30px -18px rgba(var(--shadow),.3);}
.page-head::before{content:""; position:absolute; left:0; top:0; bottom:0; width:6px;
  background:linear-gradient(180deg,var(--brand),var(--brand2));}
.page-head .pg-title {font-size:19px; font-weight:700; color:var(--ink); padding-left:8px;}
.page-head .crumb {font-size:12px; color:var(--muted);}
.page-head .crumb b {color:var(--brand-ink);}

/* ── การ์ด ───────────────────────────────────────────── */
.card {border-radius:18px !important; background:var(--card) !important;
  border:1px solid var(--line) !important;
  box-shadow:0 1px 2px rgba(var(--shadow),.04), 0 18px 40px -22px rgba(var(--shadow),.22) !important;
  padding:20px 22px !important; overflow:visible !important;
  transition:transform .2s ease, box-shadow .2s ease, border-color .2s ease;}
.card:hover {transform:translateY(-3px);
  box-shadow:0 2px 6px rgba(var(--shadow),.05), 0 28px 52px -22px rgba(var(--shadow),.32) !important;
  border-color:#d6daf2 !important;}
.section-title {font-weight:700 !important; font-size:15px !important; color:var(--ink) !important;
  position:relative; padding-left:15px !important; margin:0 0 13px !important; letter-spacing:.1px;}
.section-title::before{content:""; position:absolute; left:0; top:1px; bottom:1px; width:5px;
  border-radius:4px; background:linear-gradient(180deg,var(--brand),var(--brand2));}

/* ── ปุ่มหลัก (มีประกายเมื่อ hover) ────────────────────── */
.run-btn {position:relative; overflow:hidden;
  background:linear-gradient(120deg,var(--brand),var(--brand2) 60%, var(--brand3)) !important;
  border:none !important; color:#fff !important; font-weight:700 !important;
  font-size:16px !important; letter-spacing:.4px; border-radius:14px !important;
  padding:17px !important; margin-top:10px !important;
  box-shadow:0 14px 30px -10px rgba(var(--brand-rgb),.65) !important;
  transition:transform .15s ease, box-shadow .15s ease, filter .15s ease;}
.run-btn::after{content:""; position:absolute; top:0; left:-60%; width:42%; height:100%;
  background:linear-gradient(100deg,transparent,rgba(255,255,255,.4),transparent);
  transform:skewX(-20deg); transition:left .55s ease;}
.run-btn:hover {transform:translateY(-2px); filter:brightness(1.06) saturate(1.05);
  box-shadow:0 20px 38px -10px rgba(var(--brand-rgb),.72) !important;}
.run-btn:hover::after{left:120%;}

.ai-btn {background:#fff !important;
  border:1px solid #d8dcf3 !important; color:var(--brand-ink) !important; font-weight:600 !important;
  border-radius:12px !important; padding:12px !important; transition:all .16s ease;
  box-shadow:0 6px 16px -12px rgba(var(--brand-rgb),.6) !important;}
.ai-btn:hover {background:linear-gradient(120deg,var(--tint),var(--tint2)) !important;
  border-color:var(--brand) !important; color:var(--brand-deep) !important;
  transform:translateY(-2px); box-shadow:0 12px 24px -12px rgba(var(--brand-rgb),.7) !important;}

.status-box {border-radius:16px !important; padding:16px 18px !important; color:var(--ink) !important;
  background:linear-gradient(120deg,#fbfbff,#f3f4fe) !important; border:1px solid #e2e5f6 !important;
  font-size:14px !important; line-height:1.8; position:relative; padding-left:22px !important;
  box-shadow:0 10px 26px -20px rgba(var(--shadow),.3);}
.status-box::before{content:""; position:absolute; left:0; top:0; bottom:0; width:6px;
  border-radius:16px 0 0 16px; background:linear-gradient(180deg,var(--brand),var(--brand2));}
.status-box b {color:var(--brand-ink);}

/* ── การ์ดต้อนรับแบบคู่มือ 3 ขั้นตอน (empty state ฝั่งผลลัพธ์) ── */
.welcome-guide {padding:18px 20px 18px 22px !important;}
.welcome-guide .wg-title {font-weight:700; font-size:14.5px; color:var(--brand-deep);
  margin-bottom:12px;}
.welcome-guide .wg-step {display:flex; align-items:center; gap:12px; font-size:13.5px;
  color:var(--ink); padding:7px 0;}
.welcome-guide .wg-step b {flex:0 0 26px; width:26px; height:26px; border-radius:50%;
  display:inline-flex; align-items:center; justify-content:center; font-size:13px; color:#fff;
  background:linear-gradient(135deg,var(--brand),var(--brand2));
  box-shadow:0 6px 14px -6px rgba(var(--brand-rgb),.6);}

.hint {font-size:12.5px !important; color:var(--muted) !important; margin:5px 2px 2px !important;}
.quick-btn {border-radius:999px !important; font-size:12.5px !important; font-weight:600 !important;
  background:var(--tint) !important; border:1px solid #e0e3f5 !important; color:var(--brand-ink) !important;
  min-width:0 !important; padding:8px 15px !important; transition:all .16s ease;}
.quick-btn:hover {background:linear-gradient(120deg,var(--brand),var(--brand2)) !important;
  border-color:transparent !important; color:#fff !important; transform:translateY(-2px);
  box-shadow:0 10px 20px -8px rgba(var(--brand-rgb),.6) !important;}

/* ── โซน "ชุดสำเร็จรูป" — กล่องแยกให้ดูเรียบร้อย ────────── */
.quick-panel {border:1px solid #e2e6f6 !important; border-radius:14px !important;
  background:linear-gradient(160deg,#fbfbff,#f1f2fe) !important;
  padding:14px 15px !important; margin:6px 0 16px !important;
  box-shadow:inset 0 1px 0 rgba(255,255,255,.7) !important;}
.quick-head {font-weight:700 !important; font-size:12.5px !important;
  color:var(--brand-ink) !important; margin:0 2px 12px !important; letter-spacing:.2px;}
.quick-row {gap:8px !important; flex-wrap:wrap !important;}
.quick-panel .quick-btn {background:#ffffff !important; border-color:#dfe3f6 !important;}
.quick-panel .quick-btn:hover {background:linear-gradient(120deg,var(--brand),var(--brand2)) !important;
  border-color:transparent !important; color:#fff !important;}

/* ── เส้นคั่นหัวข้อ "เลือกตามหมวดหมู่" ───────────────────── */
.group-divider {display:flex !important; align-items:center; gap:11px;
  font-size:11.5px !important; font-weight:700 !important; color:var(--faint) !important;
  letter-spacing:.6px; margin:4px 2px 8px !important;}
.group-divider::after {content:""; flex:1; height:1px;
  background:linear-gradient(90deg,#dde1f5,transparent);}

/* ── หัวข้อหมวด + ตาราง grid ของรายการวัตถุ ───────────── */
.preset-group div[data-testid="checkbox-group"] {
  display:grid !important; grid-template-columns:repeat(auto-fit, minmax(126px,1fr)) !important; gap:6px !important;}
.preset-group div[data-testid="checkbox-group"] label {
  margin:0 !important; width:100% !important; box-sizing:border-box !important;
  border:1px solid #e3e6f5 !important; border-radius:10px !important;
  background:#fbfcff !important; padding:6px 10px !important; font-size:12.5px !important;
  line-height:1.25 !important; transition:all .15s ease;}
.preset-group div[data-testid="checkbox-group"] label:hover {
  border-color:#b6bdec !important; background:var(--tint) !important; transform:translateY(-1px);}
.preset-group div[data-testid="checkbox-group"] label:has(input:checked) {
  border-color:var(--brand) !important; background:linear-gradient(120deg,#eef0fe,#f1ecfe) !important;
  font-weight:600 !important; color:var(--brand-deep) !important;
  box-shadow:0 0 0 1px var(--brand), 0 8px 18px -10px rgba(var(--brand-rgb),.65) !important;}

/* ── dropdown รายการสัตว์: ตัวกรอง + ป๊อปอัปเป็นระเบียบ ── */
.animal-dd, .animal-dd * {overflow:visible !important;}
.animal-dd .options, .options {z-index:80 !important;}
.animal-dd .options {max-height:240px !important; overflow-y:auto !important;
  border-radius:12px !important; box-shadow:0 14px 30px -10px rgba(var(--shadow),.32) !important;}
.animal-dd .options li, .animal-dd .options .item {font-size:13px !important; padding:7px 12px !important;}
.animal-letter {margin-bottom:8px !important;}
.animal-letter .options {max-height:280px !important; overflow-y:auto !important;}

/* ── แท็บระบบ (เมนูหลัก) — segmented pill bar ───────────── */
.gov-tabs {background:transparent !important; border:none !important; box-shadow:none !important;
  margin:18px 0 8px !important;}
.gov-tabs .tab-nav, .gov-tabs [role="tablist"] {display:flex !important; gap:6px !important;
  background:#fff !important; border:1px solid var(--line) !important; border-bottom:1px solid var(--line) !important;
  padding:6px !important; border-radius:16px !important; width:max-content !important; max-width:100% !important;
  box-shadow:0 12px 30px -20px rgba(var(--shadow),.35) !important;}
.gov-tabs button[role="tab"] {color:var(--muted) !important; font-weight:600 !important;
  font-size:14px !important; background:transparent !important; border:none !important;
  border-radius:11px !important; padding:11px 22px !important; margin:0 !important;
  transition:all .16s ease;}
.gov-tabs button[role="tab"]:hover {color:var(--brand-ink) !important; background:var(--tint) !important;}
.gov-tabs button[role="tab"].selected {color:#fff !important;
  background:linear-gradient(135deg,var(--brand),var(--brand2)) !important;
  box-shadow:0 10px 22px -10px rgba(var(--brand-rgb),.65) !important;}

/* ── การ์ด KPI + เอกสารรายงาน ─────────────────────────── */
.rpt-doc {background:#fff; border:1px solid var(--line); border-radius:20px; padding:26px 28px;
  box-shadow:0 18px 44px -22px rgba(var(--shadow),.26); position:relative; overflow:hidden;}
.rpt-doc::before{content:""; position:absolute; top:0; left:0; right:0; height:5px;
  background:linear-gradient(90deg,var(--brand),var(--brand2),var(--brand3));}
.rpt-head {display:flex; align-items:center; gap:16px; border-bottom:1px solid var(--line);
  padding-bottom:18px; margin-bottom:18px;}
.rpt-emblem {width:58px; height:58px; flex:0 0 58px; border-radius:16px; font-size:28px;
  display:flex; align-items:center; justify-content:center;
  background:linear-gradient(160deg,#eef0fe,#efe9fe); border:1px solid #d9ddf5;}
.rpt-org {font-size:12px; letter-spacing:4px; color:var(--brand-ink); font-weight:800;}
.rpt-title {font-size:20px; font-weight:800; color:var(--ink); margin-top:3px;}
.rpt-meta {display:flex; flex-wrap:wrap; gap:8px 28px; font-size:12.5px; color:var(--muted);
  margin-bottom:18px;}
.rpt-meta b {color:var(--brand-ink); font-weight:600;}
.kpi-row {display:grid; grid-template-columns:repeat(auto-fit, minmax(155px, 1fr)); gap:14px;}
.kpi {position:relative; overflow:hidden; background:linear-gradient(160deg,#fbfcff,#f1f3fe);
  border:1px solid #e4e7f6; border-radius:16px; padding:20px 14px; text-align:center;
  transition:transform .16s ease, box-shadow .16s ease;}
.kpi::before{content:""; position:absolute; top:0; left:0; right:0; height:3px;
  background:linear-gradient(90deg,var(--brand),var(--brand2));}
.kpi:hover {transform:translateY(-3px); box-shadow:0 18px 32px -16px rgba(var(--brand-rgb),.4);}
.kpi-v {font-size:31px; font-weight:800; line-height:1.1;
  background:linear-gradient(120deg,var(--brand-ink),var(--brand3)); -webkit-background-clip:text;
  background-clip:text; -webkit-text-fill-color:transparent;}
.kpi-l {font-size:11.5px; color:var(--muted); margin-top:7px;}

/* ── ส่วนท้าย (footer หลายคอลัมน์) ─────────────────────── */
.app-footer {margin-top:22px; border-radius:20px; overflow:hidden;
  border:1px solid var(--line); background:#fff;
  box-shadow:0 18px 44px -26px rgba(var(--shadow),.3);}
.app-footer .ft-cols {display:flex; flex-wrap:wrap; gap:26px; padding:24px 30px;
  background:linear-gradient(180deg,#f8f9fe,#fff);}
.app-footer .ft-col {flex:1 1 200px; font-size:12.5px; color:var(--muted); line-height:2;}
.app-footer .ft-h {font-size:13px; font-weight:700; color:var(--ink);
  border-bottom:2px solid var(--brand); padding-bottom:6px; margin-bottom:10px; display:inline-block;}
.app-footer .ft-base {background:linear-gradient(90deg,var(--brand-night),#241f57); color:#c2c8ee;
  font-size:11.5px; text-align:center; padding:13px 16px; letter-spacing:.4px;}

/* ── กล่องอัปโหลดวิดีโอ (dropzone โมเดิร์น) ─────────────── */
.video-in {border-radius:16px !important; overflow:hidden !important;}
.video-in .wrap {border:2px dashed #c3c9f0 !important; border-radius:16px !important;
  min-height:172px !important; color:var(--brand-ink) !important;
  background:
    radial-gradient(460px 150px at 50% 0%, rgba(139,92,246,.10), transparent 70%),
    linear-gradient(160deg,#fbfbff,#eef1fd) !important;
  transition:border-color .2s ease, background .2s ease, box-shadow .2s ease, transform .2s ease;}
.video-in .wrap:hover {border-color:var(--brand) !important;
  background:
    radial-gradient(460px 150px at 50% 0%, rgba(139,92,246,.16), transparent 70%),
    linear-gradient(160deg,#f2f3fe,#e9ecfd) !important;
  box-shadow:inset 0 0 0 3px rgba(var(--brand-rgb),.07),
             0 16px 32px -20px rgba(var(--brand-rgb),.45) !important;
  transform:translateY(-1px);}
.video-in .wrap svg {color:var(--brand) !important; opacity:1 !important;
  width:40px !important; height:40px !important; transform:scale(1) !important;
  filter:drop-shadow(0 6px 12px rgba(var(--brand-rgb),.35));}

/* ── ป้ายใต้กล่องอัปโหลด: คำอธิบาย + ชนิดไฟล์ ──────────── */
.upload-hint {display:flex !important; align-items:center; justify-content:space-between;
  gap:10px; flex-wrap:wrap; margin-top:11px !important;}
.upload-hint .uh-text {font-size:12.5px; color:var(--muted);}
.upload-hint .uh-exts {display:inline-flex; gap:6px;}
.upload-hint .ext {font-size:10.5px; font-weight:700; letter-spacing:.6px; color:var(--brand-ink);
  background:var(--tint); border:1px solid #dde1f5; padding:3px 9px; border-radius:7px;}

/* ── การ์ดเลือกวัตถุ: หัวข้อหมวดเป็นชิป ─────────────────── */
.hint-head {font-weight:600 !important; color:var(--brand-ink) !important;
  font-size:13px !important; margin:2px 2px 10px !important;}
.cat-title {margin:11px 0 6px !important;}
.cat-title p, .cat-title span {display:inline-block !important;
  background:linear-gradient(120deg,var(--tint),var(--tint2)) !important; border:1px solid #dde1f5 !important;
  color:var(--brand-ink) !important; font-weight:700 !important; font-size:11.5px !important;
  padding:4px 12px !important; border-radius:999px !important; letter-spacing:.2px;}
"""
CSS = CSS.replace("%DOT%", _dot)

HERO = f"""
<div class="app-ubar">
  <div class="ub-left">ระบบสารสนเทศองค์กร · AIDC Tech</div>
  <div class="ub-right"><span class="ub-dot"></span> {_badge}</div>
</div>
<div class="app-masthead">
  <span class="mh-aurora a1"></span>
  <span class="mh-aurora a2"></span>
  <span class="mh-aurora a3"></span>
  <div class="mh-emblem">🛡️</div>
  <div class="mh-titles">
    <div class="org">AIDC TECH</div>
    <h1>ລະບົບປະມວນຜົນ ແລະ ວິເຄາะວິດີໂອດ້ວຍປັນຍາປະດິດ</h1>
    <div class="sub">AIDC Tech Video Processor — AI Video Analytics System</div>
  </div>
  <div class="mh-ver"><b>เวอร์ชัน 1.0</b><br>ระบบวิเคราะห์อัจฉริยะ</div>
</div>
"""

PAGE_HEAD = """
<div class="page-head">
  <div class="pg-title">การวิเคราะห์วิดีโอด้วยปัญญาประดิษฐ์</div>
  <div class="crumb">หน้าหลัก › ระบบวิเคราะห์วิดีโอ › <b>ประมวลผล</b></div>
</div>
"""

FOOTER = """
<div class="app-footer">
  <div class="ft-cols">
    <div class="ft-col">
      <div class="ft-h">AIDC Tech</div>
      ระบบวิเคราะห์วิดีโออัจฉริยะ<br>สำหรับหน่วยงานภาครัฐ
    </div>
    <div class="ft-col">
      <div class="ft-h">เกี่ยวกับระบบ</div>
      เวอร์ชัน 1.0<br>ขับเคลื่อนด้วย Ultralytics YOLO และ Claude AI
    </div>
    <div class="ft-col">
      <div class="ft-h">การสนับสนุน</div>
      ติดต่อผู้ดูแลระบบ<br>คู่มือและเอกสารการใช้งาน
    </div>
  </div>
  <div class="ft-base">สงวนลิขสิทธิ์ &copy; 2569 AIDC Tech · เพื่อการใช้งานภายในหน่วยงาน</div>
</div>
"""

WELCOME = (
    '<div class="status-box welcome-guide">'
    '<div class="wg-title">เริ่มต้นใช้งานใน 3 ขั้นตอน</div>'
    '<div class="wg-step"><b>1</b><span>อัปโหลดไฟล์วิดีโอทางด้านซ้าย</span></div>'
    '<div class="wg-step"><b>2</b><span>เลือกวัตถุที่ต้องการตรวจจับ</span></div>'
    '<div class="wg-step"><b>3</b><span>กดปุ่ม “เริ่มประมวลผล” — ผลการวิเคราะห์จะแสดงตรงนี้</span></div>'
    '</div>'
)

with gr.Blocks(title="AIDC Tech Video Processor", fill_width=False) as demo:
    gr.HTML(HERO)

    report_state = gr.State(None)      # ผลตรวจจับล่าสุด (แชร์ข้ามแท็บ)
    ai_report_state = gr.State("")     # ข้อความรายงาน AI ล่าสุด

    with gr.Tabs(elem_classes="gov-tabs"):
        # ═════════ แท็บที่ 1: วิเคราะห์วิดีโอ ═════════
        with gr.Tab("วิเคราะห์วิดีโอ"):
            gr.HTML(PAGE_HEAD)
            with gr.Row(equal_height=False):
                # ───────── ฝั่งซ้าย: ตั้งค่า ─────────
                with gr.Column(scale=4):
                    with gr.Group(elem_classes="card"):
                        gr.Markdown("ขั้นที่ 1 · อัปโหลดวิดีโอ", elem_classes="section-title")
                        video_in = gr.Video(label="", height=210, elem_classes="video-in")
                        gr.HTML(
                            '<div class="upload-hint">'
                            '<span class="uh-text">ลากไฟล์มาวาง หรือคลิกที่กรอบเพื่อเลือก</span>'
                            '<span class="uh-exts">'
                            '<span class="ext">MP4</span><span class="ext">MOV</span>'
                            '<span class="ext">AVI</span></span>'
                            '</div>')

                    with gr.Group(elem_classes="card obj-card"):
                        gr.Markdown("ขั้นที่ 2 · เลือกวัตถุที่จะตรวจจับ",
                                    elem_classes="section-title")
                        # โซนเลือกเร็ว — จัดเป็นกล่องแยกให้ดูเป็นระเบียบ
                        with gr.Group(elem_classes="quick-panel"):
                            gr.Markdown("⚡ ชุดสำเร็จรูป · กดเลือกทีเดียวหลายอย่าง",
                                        elem_classes="quick-head")
                            with gr.Row(elem_classes="quick-row"):
                                quick_btns = [gr.Button(name, size="sm", elem_classes="quick-btn")
                                              for name in QUICK_SETS]
                        gr.Markdown("เลือกตามหมวดหมู่", elem_classes="group-divider")
                        # CheckboxGroup แยกตามหมวด เรียงเป็นตาราง grid 2 คอลัมน์
                        preset_groups = []
                        for cat_label, items in CATEGORIES.items():
                            gr.Markdown(cat_label, elem_classes="cat-title")
                            cg = gr.CheckboxGroup(
                                choices=list(items.keys()), value=[], show_label=False,
                                container=False, elem_classes="preset-group")
                            preset_groups.append(cg)

                        gr.Markdown(f"สัตว์ (รายการละเอียด {len(ANIMALS)} ชนิด)",
                                    elem_classes="cat-title")
                        animal_letter = gr.Dropdown(
                            choices=ANIMAL_LETTER_CHOICES, value=ANIMAL_ALL,
                            label="กรองตามตัวอักษร (A–Z)", elem_classes="animal-letter")
                        animal_dd = gr.Dropdown(
                            choices=ANIMALS, value=[], multiselect=True, filterable=True,
                            show_label=False, elem_classes="animal-dd",
                            info="ย่อรายการด้วยตัวกรองด้านบน หรือพิมพ์ค้นหาชื่อสัตว์ (อังกฤษ) "
                                 "ได้เลย — รายการที่เลือกไว้จะไม่หายเมื่อสลับหมวด")

                        custom_tb = gr.Textbox(
                            label="เพิ่มเอง (พิมพ์ภาษาอังกฤษ คั่นด้วย ,)",
                            placeholder="เช่น: bridge, waterfall, helicopter",
                            info="ติ๊กหรือพิมพ์อย่างใดอย่างหนึ่ง = ใช้ AI ตรวจจับตามนั้น | "
                                 "ไม่เลือกอะไรเลย = ใช้โมเดลปกติ 80 ชนิด",
                        )
                        model_dd = gr.Dropdown(
                            find_models(), value=find_models()[0],
                            label="โมเดล .pt (ใช้เมื่อไม่ได้เลือก/พิมพ์อะไร)")

                    with gr.Accordion("ตั้งค่าขั้นสูง", open=False, elem_classes="card"):
                        with gr.Row():
                            conf_sl = gr.Slider(0.05, 0.95, value=0.25, step=0.05,
                                                label="ความมั่นใจขั้นต่ำ (confidence)")
                            stride_sl = gr.Slider(1, 10, value=1, step=1,
                                                  label="ทุก ๆ N เฟรม (เพิ่ม = เร็วขึ้น)")
                        gpu_ck = gr.Checkbox(value=HAS_CUDA, label="ใช้ GPU เร่งความเร็ว",
                                             interactive=HAS_CUDA)

                    with gr.Group(elem_classes="card"):
                        gr.Markdown("ตัวเลือกเสริม (segmentation / นับจำนวน)",
                                    elem_classes="section-title")
                        with gr.Row():
                            seg_ck = gr.Checkbox(
                                value=False,
                                label="🎭 แสดงขอบเขตแบบ mask",
                                info="ระบายพื้นที่วัตถุแทนกรอบ · ใช้กับ 80 ชนิดมาตรฐาน")
                            track_ck = gr.Checkbox(
                                value=False,
                                label="🔢 นับจำนวนตัวจริง",
                                info="ติดตามด้วย ID นับตัวไม่ซ้ำ · แนะนำตั้ง N เฟรม = 1")

                    run_btn = gr.Button("ขั้นที่ 3 · ▶  เริ่มประมวลผล", elem_classes="run-btn")

                # ───────── ฝั่งขวา: ผลลัพธ์ ─────────
                with gr.Column(scale=6):
                    status = gr.HTML(WELCOME)
                    with gr.Group(elem_classes="card"):
                        gr.Markdown("วิดีโอผลการประมวลผล", elem_classes="section-title")
                        video_out = gr.Video(label="", height=320)
                    with gr.Row(equal_height=False):
                        with gr.Column(scale=1):
                            with gr.Group(elem_classes="card"):
                                gr.Markdown("สรุปผลการตรวจจับ", elem_classes="section-title")
                                table_out = gr.Dataframe(label="", interactive=False, wrap=True)
                        with gr.Column(scale=1):
                            with gr.Group(elem_classes="card"):
                                gr.Markdown("รายงานวิเคราะห์ด้วย AI", elem_classes="section-title")
                                report_btn = gr.Button("📝 สร้างรายงานประเมินความเสี่ยง (Claude AI)",
                                                       elem_classes="ai-btn")
                                report_md = gr.Markdown("")

        # ═════════ แท็บที่ 2: รายงานผล ═════════
        with gr.Tab("รายงานผล") as report_tab:
            gr.HTML('<div class="page-head"><div class="pg-title">รายงานผลการวิเคราะห์</div>'
                    '<div class="crumb">หน้าหลัก › <b>รายงานผล</b></div></div>')
            with gr.Row():
                load_report_btn = gr.Button("🔄 โหลดผลการวิเคราะห์ล่าสุด",
                                            elem_classes="ai-btn")
            report_page_html = gr.HTML(
                '<div class="status-box">กดปุ่ม “โหลดผลการวิเคราะห์ล่าสุด” '
                'เพื่อแสดงรายงาน (หรือสลับมาที่แท็บนี้หลังประมวลผลเสร็จ)</div>')
            with gr.Group(elem_classes="card"):
                gr.Markdown("ภาพเฟรมตัวอย่าง (ตรวจพบวัตถุมากที่สุด)",
                            elem_classes="section-title")
                report_keyframe_img = gr.Image(label="", height=320, show_label=False)
            with gr.Group(elem_classes="card"):
                gr.Markdown("ตารางสรุปการตรวจจับ", elem_classes="section-title")
                report_table = gr.Dataframe(label="", interactive=False, wrap=True)
            with gr.Group(elem_classes="card"):
                gr.Markdown("รายงานประเมินความเสี่ยง (วิเคราะห์โดย AI)",
                            elem_classes="section-title")
                report_ai_md = gr.Markdown("")
            with gr.Group(elem_classes="card"):
                gr.Markdown("ส่งออกเอกสารและไฟล์", elem_classes="section-title")
                with gr.Row():
                    word_btn = gr.Button("📄 สร้างรายงาน Word", elem_classes="ai-btn")
                    excel_btn = gr.Button("📊 สร้างตาราง Excel", elem_classes="ai-btn")
                    video_btn = gr.Button("🎬 เตรียมวิดีโอผล", elem_classes="ai-btn")
                word_file = gr.File(label="รายงาน (Word)", interactive=False)
                excel_file = gr.File(label="ตาราง (Excel)", interactive=False)
                video_file = gr.File(label="วิดีโอผลการประมวลผล", interactive=False)

    gr.HTML(FOOTER)

    # ปุ่มชุดสำเร็จรูป → เพิ่มรายการเข้าหมวดที่เกี่ยวข้อง (รวมกับที่เลือกไว้เดิม)
    def _add_set(to_add):
        def fn(*current_vals):
            updates = []
            for items, vals in zip(CATEGORIES.values(), current_vals):
                add_here = [x for x in to_add if x in items]
                updates.append(gr.update(
                    value=list(dict.fromkeys((vals or []) + add_here))))
            return updates
        return fn

    for btn, name in zip(quick_btns, QUICK_SETS):
        btn.click(_add_set(QUICK_SETS[name]),
                  inputs=preset_groups, outputs=preset_groups)

    # ตัวกรองตามตัวอักษร → ย่อรายการในดรอปดาวน์สัตว์ (คงรายการที่เลือกไว้)
    animal_letter.change(filter_animals,
                         inputs=[animal_letter, animal_dd], outputs=animal_dd)

    n_groups = len(preset_groups)

    def _run(video, model_path, *rest, progress=gr.Progress()):
        group_vals = rest[:n_groups]
        (animal_labels, custom_classes, segment, track,
         conf, use_gpu, stride) = rest[n_groups:]
        preset_labels = [x for vals in group_vals for x in (vals or [])]
        video_out_, table, summary, report_data = process_video(
            video, model_path, preset_labels, animal_labels, custom_classes,
            segment, track, conf, use_gpu, stride, progress=progress)
        # ล้างรายงาน AI เก่า + เก็บ report_data ไว้ใน State
        return (video_out_, table, f'<div class="status-box">{summary}</div>',
                report_data, "", "")

    run_btn.click(
        _run,
        inputs=[video_in, model_dd, *preset_groups, animal_dd, custom_tb,
                seg_ck, track_ck, conf_sl, gpu_ck, stride_sl],
        outputs=[video_out, table_out, status, report_state, report_md, ai_report_state],
    )

    def _gen_report(report_data, progress=gr.Progress()):
        text = generate_ai_report(report_data, progress)
        return text, text   # แสดงในแท็บวิเคราะห์ + เก็บไว้ใช้ในหน้ารายงาน

    report_btn.click(_gen_report, inputs=report_state,
                     outputs=[report_md, ai_report_state])

    # หน้า "รายงานผล": โหลดด้วยปุ่ม หรืออัตโนมัติเมื่อสลับมาที่แท็บ
    _report_outputs = [report_page_html, report_keyframe_img, report_table, report_ai_md]
    load_report_btn.click(build_report_page,
                          inputs=[report_state, ai_report_state], outputs=_report_outputs)
    report_tab.select(build_report_page,
                      inputs=[report_state, ai_report_state], outputs=_report_outputs)

    # ส่งออกเอกสาร/ไฟล์
    word_btn.click(export_word, inputs=[report_state, ai_report_state], outputs=word_file)
    excel_btn.click(export_excel, inputs=report_state, outputs=excel_file)
    video_btn.click(get_result_video, inputs=report_state, outputs=video_file)


if __name__ == "__main__":
    demo.queue().launch(inbrowser=True, theme=THEME, css=CSS)
